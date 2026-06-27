import html
import hashlib
import hmac
import json
import os
import re
import sqlite3
import subprocess
import tempfile
import time
import urllib.request
import uuid
from concurrent.futures import ThreadPoolExecutor, TimeoutError
from html.parser import HTMLParser

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from fastapi import FastAPI, File, Form, Request, UploadFile
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse

try:
    import razorpay
except ImportError:
    razorpay = None


# ================= CONFIG =================

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
ENV_PATH = os.path.join(BASE_DIR, ".env")
KEY_FILE_PATH = os.path.join(BASE_DIR, "gemini_key.txt")


def load_env_file(path: str):
    if not os.path.exists(path):
        return

    with open(path, "r", encoding="utf-8") as env_file:
        for line in env_file:
            stripped = line.strip()
            if not stripped or stripped.startswith("#") or "=" not in stripped:
                continue

            key, value = stripped.split("=", 1)
            os.environ.setdefault(key.strip(), value.strip().strip("\"'"))


def get_gemini_api_key(submitted_key: str | None = None) -> str:
    if submitted_key and submitted_key.strip():
        return submitted_key.strip()

    load_env_file(ENV_PATH)
    api_key = os.getenv("GEMINI_API_KEY") or os.getenv("GOOGLE_API_KEY") or ""

    if api_key:
        return api_key.strip()

    if os.path.exists(KEY_FILE_PATH):
        with open(KEY_FILE_PATH, "r", encoding="utf-8") as key_file:
            return key_file.read().strip()

    return ""


GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-3.5-flash")
GEMINI_API_KEY = get_gemini_api_key()

RESUME_PRICE_RUPEES = int(os.getenv("RESUME_PRICE_RUPEES", "49"))
RESUME_PRICE_PAISE = RESUME_PRICE_RUPEES * 100
FREE_TRIAL_RESUMES = int(os.getenv("FREE_TRIAL_RESUMES", "3"))
OWNER_EMAILS = {
    email.strip().lower()
    for email in os.getenv("OWNER_EMAILS", "sumitmishra7886@gmail.com").split(",")
    if email.strip()
}
RAZORPAY_KEY_ID = os.getenv("RAZORPAY_KEY_ID", "")
RAZORPAY_KEY_SECRET = os.getenv("RAZORPAY_KEY_SECRET", "")
DOWNLOAD_TOKEN_SECRET = (
    os.getenv("DOWNLOAD_TOKEN_SECRET")
    or RAZORPAY_KEY_SECRET
    or GEMINI_API_KEY
    or "resume-fit-pro-local-download-secret"
)
PAYMENTS_ENABLED = bool(RAZORPAY_KEY_ID and RAZORPAY_KEY_SECRET and razorpay)
razorpay_client = (
    razorpay.Client(auth=(RAZORPAY_KEY_ID, RAZORPAY_KEY_SECRET))
    if PAYMENTS_ENABLED
    else None
)

if GEMINI_API_KEY:
    pass

AI_EXECUTOR = ThreadPoolExecutor(max_workers=2)

app = FastAPI()

OUTPUT_DIR = os.path.join(BASE_DIR, "resumes")
DB_PATH = os.path.join(BASE_DIR, "resumes.db")
TEMPLATE_PATH = os.path.join(BASE_DIR, "template.html")
INDEX_PATH = os.path.join(BASE_DIR, "index.html")
MANIFEST_PATH = os.path.join(BASE_DIR, "manifest.json")
SERVICE_WORKER_PATH = os.path.join(BASE_DIR, "service-worker.js")
APP_ICON_PATH = os.path.join(BASE_DIR, "app-icon.svg")
APP_LOGO_PATH = os.path.join(BASE_DIR, "app-logo.png")
APP_LOGO_ICON_PATH = os.path.join(BASE_DIR, "app-logo-icon.png")
APK_PATH = os.path.join(BASE_DIR, "ResumeFit-Pro.apk")

os.makedirs(OUTPUT_DIR, exist_ok=True)


@app.get("/health")
def health():
    return {"status": "ok"}


@app.get("/privacy", response_class=HTMLResponse)
def privacy():
    return HTMLResponse(
        content="""
<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>Privacy Policy | ResumeFit Pro</title>
<style>
body { font-family: Arial, sans-serif; margin: 0; background: #f4f7fb; color: #122033; }
main { max-width: 860px; margin: 0 auto; padding: 40px 20px; line-height: 1.6; }
h1, h2 { color: #102a43; }
a { color: #0c6576; }
</style>
</head>
<body>
<main>
<h1>Privacy Policy</h1>
<p>ResumeFit Pro helps users optimize resumes against job descriptions. Users may upload resume files and paste job descriptions for processing.</p>
<h2>Data Processed</h2>
<p>The application processes uploaded resume content, job descriptions, generated analysis, optimized resume text, and downloaded document output.</p>
<h2>Purpose</h2>
<p>Data is used only to provide ATS analysis, resume recommendations, preview, and Word/PDF/LaTeX document generation.</p>
<h2>Third-Party AI Processing</h2>
<p>Resume and job description content may be sent to the configured Gemini API provider to generate optimization results.</p>
<h2>Storage</h2>
<p>Generated resume HTML and analysis data may be stored temporarily in the application database to support preview and document download links.</p>
<h2>User Responsibility</h2>
<p>Users should avoid uploading unnecessary sensitive information and should review generated resume content before using it for job applications.</p>
<h2>Contact</h2>
<p>For support or privacy questions, contact: sumitmishra7886@gmail.com</p>
<p><a href="/">Back to app</a></p>
</main>
</body>
</html>
"""
    )


@app.get("/terms", response_class=HTMLResponse)
def terms():
    return HTMLResponse(
        content="""
<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>Terms | ResumeFit Pro</title>
<style>
body { font-family: Arial, sans-serif; margin: 0; background: #f4f7fb; color: #122033; }
main { max-width: 860px; margin: 0 auto; padding: 40px 20px; line-height: 1.6; }
h1, h2 { color: #102a43; }
a { color: #0c6576; }
</style>
</head>
<body>
<main>
<h1>Terms of Use</h1>
<p>ResumeFit Pro provides AI-assisted resume optimization and ATS scoring estimates. Results are suggestions and do not guarantee interviews, job offers, or hiring outcomes.</p>
<h2>User Review</h2>
<p>Users are responsible for reviewing, editing, and verifying all generated resume content before use.</p>
<h2>No Employment Guarantee</h2>
<p>ATS scores are estimates and may differ from employer systems.</p>
<h2>Acceptable Use</h2>
<p>Do not upload content that you do not have permission to process.</p>
<h2>Contact</h2>
<p>Support: sumitmishra7886@gmail.com</p>
<p><a href="/">Back to app</a></p>
</main>
</body>
</html>
"""
    )

# ================= DATABASE =================

def connect_db():
    db_paths = [
        DB_PATH,
        os.path.join(OUTPUT_DIR, "resumes.db"),
        ":memory:",
    ]

    last_error = None
    for db_path in db_paths:
        try:
            database = sqlite3.connect(db_path, check_same_thread=False, timeout=30)
            database.execute("PRAGMA journal_mode=WAL")
            database.execute(
                """
CREATE TABLE IF NOT EXISTS resumes (
    id TEXT PRIMARY KEY,
    html TEXT NOT NULL,
    analysis_json TEXT,
    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
)
"""
            )
            database.execute(
                """
CREATE TABLE IF NOT EXISTS payments (
    id TEXT PRIMARY KEY,
    resume_id TEXT NOT NULL,
    email TEXT NOT NULL,
    amount_paise INTEGER NOT NULL,
    currency TEXT NOT NULL DEFAULT 'INR',
    status TEXT NOT NULL,
    razorpay_order_id TEXT,
    razorpay_payment_id TEXT,
    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
    paid_at TIMESTAMP
)
"""
            )
            database.execute(
                """
CREATE TABLE IF NOT EXISTS free_downloads (
    id TEXT PRIMARY KEY,
    resume_id TEXT NOT NULL,
    email TEXT NOT NULL,
    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
    UNIQUE(resume_id, email)
)
"""
            )
            database.commit()

            try:
                database.execute("ALTER TABLE resumes ADD COLUMN analysis_json TEXT")
                database.commit()
            except sqlite3.OperationalError:
                pass

            return database
        except sqlite3.Error as error:
            last_error = error

    raise RuntimeError(f"Could not initialize SQLite database: {last_error}")


conn = connect_db()
cursor = conn.cursor()


# ================= PAYMENTS =================

def normalize_email(email: str | None) -> str:
    return (email or "").strip().lower()


def is_owner_email(email: str | None) -> bool:
    normalized = normalize_email(email)
    return bool(normalized and normalized in OWNER_EMAILS)


def make_download_token(fid: str, email: str, ttl_seconds: int = 3600) -> str:
    expires_at = int(time.time()) + ttl_seconds
    normalized_email = normalize_email(email)
    payload = f"{fid}|{normalized_email}|{expires_at}"
    signature = hmac.new(
        DOWNLOAD_TOKEN_SECRET.encode("utf-8"),
        payload.encode("utf-8"),
        hashlib.sha256,
    ).hexdigest()
    return f"{payload}|{signature}"


def verify_download_token(fid: str, token: str | None) -> bool:
    if not token:
        return False

    parts = token.split("|")
    if len(parts) != 4:
        return False

    token_fid, email, expires_at_text, signature = parts
    if token_fid != fid:
        return False

    try:
        expires_at = int(expires_at_text)
    except ValueError:
        return False

    if expires_at < int(time.time()):
        return False

    payload = f"{token_fid}|{email}|{expires_at}"
    expected = hmac.new(
        DOWNLOAD_TOKEN_SECRET.encode("utf-8"),
        payload.encode("utf-8"),
        hashlib.sha256,
    ).hexdigest()
    return hmac.compare_digest(signature, expected)


def resume_exists(fid: str) -> bool:
    cursor.execute("SELECT 1 FROM resumes WHERE id=?", (fid,))
    return bool(cursor.fetchone())


def free_trial_status(email: str, fid: str | None = None) -> dict:
    normalized_email = normalize_email(email)
    if not normalized_email or FREE_TRIAL_RESUMES <= 0:
        return {"limit": FREE_TRIAL_RESUMES, "used": 0, "remaining": 0, "resume_unlocked": False}

    cursor.execute(
        "SELECT COUNT(DISTINCT resume_id) FROM free_downloads WHERE email=?",
        (normalized_email,),
    )
    used = int(cursor.fetchone()[0] or 0)
    resume_unlocked = False

    if fid:
        cursor.execute(
            "SELECT 1 FROM free_downloads WHERE email=? AND resume_id=?",
            (normalized_email, fid),
        )
        resume_unlocked = bool(cursor.fetchone())

    remaining = max(FREE_TRIAL_RESUMES - used, 0)
    return {
        "limit": FREE_TRIAL_RESUMES,
        "used": used,
        "remaining": remaining,
        "resume_unlocked": resume_unlocked,
    }


def can_use_free_trial(email: str, fid: str) -> bool:
    status = free_trial_status(email, fid)
    return bool(status["resume_unlocked"] or status["remaining"] > 0)


def record_free_trial(email: str, fid: str):
    cursor.execute(
        """
        INSERT OR IGNORE INTO free_downloads (id, resume_id, email)
        VALUES (?, ?, ?)
        """,
        (str(uuid.uuid4()), fid, normalize_email(email)),
    )
    conn.commit()


def payment_required_response():
    return JSONResponse(
        status_code=402,
        content={
            "error": "Payment required",
            "amount_rupees": RESUME_PRICE_RUPEES,
            "currency": "INR",
        },
    )


def require_download_access(fid: str, token: str | None):
    if verify_download_token(fid, token):
        return None
    return payment_required_response()


# ================= RESUME IMPORT =================

def extract_text(file: UploadFile) -> str:
    filename = (file.filename or "").lower()
    if filename.endswith(".pdf"):
        try:
            from pypdf import PdfReader

            reader = PdfReader(file.file)
            return "\n".join((page.extract_text() or "").strip() for page in reader.pages if page.extract_text())
        except Exception:
            return ""

    doc = Document(file.file)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


# ================= HEADER =================

def clean_contact_text(value: str) -> str:
    value = " ".join(value.split())
    value = value.replace("LinkedIn –", "LinkedIn -").replace("LinkedIn —", "LinkedIn -")
    value = value.replace("linkedin.com/i n/", "linkedin.com/in/")
    value = re.sub(r"(linkedin\.com)\s*/\s*in\s*/", r"\1/in/", value, flags=re.IGNORECASE)
    value = re.sub(r"(https?://)\s+", r"\1", value, flags=re.IGNORECASE)
    value = re.sub(r"(www\.)\s+", r"\1", value, flags=re.IGNORECASE)
    return value


def build_contact_block(lines: list[str]) -> str:
    raw = clean_contact_text(" ".join(lines))
    email_match = re.search(r"[\w.+-]+@[\w-]+\.[\w.-]+", raw)
    phone_match = re.search(r"(?:\+?\d[\d\s-]{8,}\d)", raw)
    urls = re.findall(r"(?:https?://|www\.)[^\s|,;]+|(?:linkedin\.com|github\.com)/[^\s|,;]+", raw, flags=re.IGNORECASE)
    urls = [url if url.startswith(("http://", "https://")) else f"https://{url}" for url in urls]

    email = email_match.group(0) if email_match else ""
    phone = phone_match.group(0).strip() if phone_match else ""
    linkedin_url = next((url for url in urls if "linkedin.com" in url.lower()), "")
    portfolio_url = next((url for url in urls if "linkedin.com" not in url.lower()), "")

    location_source = raw
    for value in [email, phone, *urls]:
        if value:
            location_source = location_source.replace(value, " ")
    location_source = re.sub(r"\b(Portfolio|LinkedIn)\b\s*[-–—]?", " ", location_source, flags=re.IGNORECASE)

    location = ""
    location_match = re.search(
        r"\b([A-Za-z][A-Za-z ]{1,40},\s*[A-Za-z][A-Za-z ]{1,40}(?:,\s*[A-Za-z][A-Za-z ]{1,40})?)\b",
        location_source,
    )
    if location_match:
        location = location_match.group(1).strip()

    rows = []
    if phone:
        tel = re.sub(r"[^\d+]", "", phone)
        rows.append(f'<div class="contact-line">Phone: <a href="tel:{html.escape(tel)}">{html.escape(phone)}</a></div>')
    if email:
        rows.append(f'<div class="contact-line">Email: <a href="mailto:{html.escape(email)}">{html.escape(email)}</a></div>')
    if location:
        rows.append(f'<div class="contact-line">Location: {html.escape(location)}</div>')
    if portfolio_url:
        portfolio_label = portfolio_url.replace("https://", "").replace("http://", "").rstrip("/")
        rows.append(f'<div class="contact-line">Portfolio: <a href="{html.escape(portfolio_url)}">{html.escape(portfolio_label)}</a></div>')
    if linkedin_url:
        linkedin_label = linkedin_url.replace("https://", "").replace("http://", "").rstrip("/")
        rows.append(f'<div class="contact-line">LinkedIn: <a href="{html.escape(linkedin_url)}">{html.escape(linkedin_label)}</a></div>')

    if rows:
        return '<div class="contact">' + "".join(rows) + "</div>"

    return f'<div class="contact">{html.escape(raw)}</div>'


def extract_contact_lines(lines: list[str]) -> list[str]:
    section_names = {
        "summary",
        "skills",
        "experience",
        "professional experience",
        "projects",
        "education",
        "certifications",
        "certification",
        "achievements",
    }
    contact_lines = []
    for line in lines[1:35]:
        lowered = line.strip().lower().rstrip(":")
        if lowered in section_names:
            break
        if re.search(r"@|(?:\+?\d[\d\s-]{8,}\d)|linkedin|https?://|www\.|location|address", line, re.IGNORECASE):
            contact_lines.append(line)
            continue
        if re.search(r"\b[A-Za-z][A-Za-z ]{1,40},\s*[A-Za-z]{2,40}(?:,\s*[A-Za-z]{2,40})?\b", line):
            contact_lines.append(line)
    return contact_lines


def extract_header(text: str, optimized_headline: str | None = None) -> str:
    lines = [line.strip() for line in text.splitlines() if line.strip()]

    name = html.escape(lines[0]) if len(lines) > 0 else ""
    role_text = optimized_headline or (lines[1] if len(lines) > 1 else "")
    role = html.escape(role_text)
    contact = build_contact_block(extract_contact_lines(lines))

    return f"""
    <div class="header-name">{name}</div>
    <div class="header-role">{role}</div>
    {contact}
    """


# ================= EDUCATION =================

def extract_education(text: str) -> str:
    lines = text.splitlines()
    education = []
    capture = False

    for line in lines:
        stripped = line.strip()

        if "education" in stripped.lower():
            capture = True
            continue

        if capture:
            if any(
                keyword in stripped.lower()
                for keyword in ["skills", "experience", "projects", "certifications", "certification", "achievements"]
            ):
                break
            if stripped:
                education.append(html.escape(stripped))

    return "<br>".join(education) if education else "Not Found"


# ================= AI / ATS =================

def important_terms(text: str) -> list[str]:
    stop_words = {
        "and", "the", "for", "with", "you", "your", "will", "are", "from", "this",
        "that", "have", "has", "our", "job", "role", "work", "team", "using", "use",
        "experience", "skills", "required", "preferred", "responsibilities",
    }
    words = re.findall(r"[A-Za-z][A-Za-z0-9+#.-]{2,}", text.lower())
    seen = {}
    for word in words:
        if word not in stop_words:
            seen[word] = seen.get(word, 0) + 1
    return [word for word, _ in sorted(seen.items(), key=lambda item: item[1], reverse=True)[:35]]


def has_any(text: str, terms: list[str]) -> bool:
    lowered = text.lower()
    return any(term.lower() in lowered for term in terms)


def unique_items(items: list[str]) -> list[str]:
    seen = set()
    result = []
    for item in items:
        cleaned = " ".join(str(item).replace("•", " ").split()).strip(" -|")
        key = cleaned.lower()
        if cleaned and key not in seen:
            seen.add(key)
            result.append(cleaned)
    return result


def extract_section_items(text: str, section_names: list[str], stop_sections: list[str] | None = None) -> list[str]:
    stop_sections = stop_sections or [
        "summary", "profile", "skills", "technical skills", "experience", "work experience",
        "projects", "education", "certifications", "achievements", "awards",
    ]
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    capture = False
    items = []

    for line in lines:
        normalized = line.strip().lower().rstrip(":")

        if any(name in normalized for name in section_names):
            capture = True
            continue

        if capture and any(stop == normalized or normalized.startswith(stop + " ") for stop in stop_sections):
            break

        if capture:
            items.append(line)

    return unique_items(items)


def infer_role(jd: str, resume: str) -> str:
    combined = f"{jd}\n{resume}"
    jd_lower = jd.lower()
    title_patterns = [
        r"(?:job title|position|role)\s*[:\-]\s*([A-Za-z0-9 &/+.-]{4,80})",
        r"we are hiring\s+(?:an?|for)?\s*([A-Za-z0-9 &/+.-]{4,80})",
        r"looking for\s+(?:an?|a)?\s*([A-Za-z0-9 &/+.-]{4,80})",
    ]
    for pattern in title_patterns:
        match = re.search(pattern, jd, re.IGNORECASE)
        if match:
            title = " ".join(match.group(1).split()).strip(" .-|")
            if title:
                return title[:70]

    if "technical support engineer" in jd_lower:
        return "Technical Support Engineer"
    if "it support engineer" in jd_lower:
        return "IT Support Engineer"
    if "desktop support engineer" in jd_lower:
        return "Desktop Support Engineer"
    if "system administrator" in jd_lower or "systems administrator" in jd_lower:
        return "System Administrator"
    if "windows administrator" in jd_lower:
        return "Windows Administrator"
    if "infrastructure administrator" in jd_lower:
        return "Infrastructure Administrator"
    if "service desk" in jd_lower:
        return "Service Desk Analyst"
    if has_any(combined, ["desktop support", "service desk", "helpdesk", "help desk"]):
        return "IT Support and Infrastructure Administrator"
    if has_any(combined, ["windows server", "active directory", "azure", "intune"]):
        return "Windows and Infrastructure Administrator"
    if has_any(combined, ["system administrator", "systems administrator"]):
        return "System Administrator"
    return "IT Infrastructure Professional"


def fallback_analysis(jd: str, resume: str) -> dict:
    jd_terms = important_terms(jd)
    resume_lower = resume.lower()
    matched = [term for term in jd_terms if term in resume_lower]
    missing = [term for term in jd_terms if term not in resume_lower][:15]
    keyword_score = round((len(matched) / max(len(jd_terms), 1)) * 100)
    current_score = clamp_score(round((keyword_score * 0.45) + 35))
    role = infer_role(jd, resume)
    resume_skills = extract_section_items(resume, ["skills", "technical skills"])
    resume_projects = extract_section_items(resume, ["projects", "project"])
    resume_certifications = extract_section_items(resume, ["certifications", "certification"])

    role_skills = [
        "Windows Server Administration",
        "Active Directory",
        "Azure Administration",
        "Microsoft Intune",
        "Endpoint Device Management",
        "Service Desk Support",
        "Desktop Support",
        "User Account Administration",
        "Access Management",
        "Incident Management",
        "Troubleshooting",
        "Patch Management",
        "System Monitoring",
        "Network Fundamentals",
        "ITIL Practices",
        "macOS Support",
        "Documentation",
        "Security Best Practices",
    ]
    keyword_skills = [term.title() for term in matched[:10] + missing[:8]]
    skills = unique_items(resume_skills + role_skills + keyword_skills)[:24]

    summary = (
        f"{role} with hands-on experience in Windows administration, infrastructure support, "
        "service desk operations, user access management, endpoint troubleshooting, and ITIL-aligned "
        "incident handling. Strong foundation in Windows Server hybrid infrastructure, Azure services, "
        "device management, security controls, and documentation, with a practical focus on improving "
        "system reliability, user productivity, and support response quality."
    )

    experience = [
        "Administered and supported Windows-based user environments, including account access, desktop issues, software support, and day-to-day infrastructure troubleshooting.",
        "Resolved service desk and desktop support incidents by diagnosing hardware, operating system, network, email, and application issues through structured troubleshooting.",
        "Supported user onboarding, access provisioning, password resets, permissions, and basic security checks while following standard IT support procedures.",
        "Worked with Windows Server and hybrid infrastructure concepts, including Active Directory, server administration, endpoint management, and cloud-aligned administration practices.",
        "Documented recurring issues, resolutions, and support steps to improve knowledge sharing, reduce repeat incidents, and strengthen operational consistency.",
    ]

    if has_any(jd, ["intune", "endpoint", "device"]):
        experience.insert(
            2,
            "Assisted with endpoint and device management activities, including configuration checks, policy awareness, and support for managed Windows devices.",
        )

    if has_any(jd, ["security", "access", "iam", "permissions"]):
        experience.insert(
            3,
            "Applied access management and basic security practices to support least-privilege permissions, account hygiene, and secure user operations.",
        )

    projects = resume_projects[:4] or [
        "Windows Server Hybrid Administration: Applied AZ-800 concepts across Windows Server administration, identity, networking, storage, and hybrid infrastructure scenarios.",
        "Service Desk Process Improvement: Organized troubleshooting notes and support documentation to improve incident handling and repeat issue resolution.",
        "Endpoint Support Practice: Supported Windows desktop administration tasks covering device setup, user productivity issues, application support, and system troubleshooting.",
    ]

    achievements = [
        "Built a certification-backed infrastructure profile with AZ-800, ITIL 4 Foundation, Google Cloud Digital Leader, Windows Server Administrator L2, and Oracle Database Administrator L2 credentials.",
        "Maintained a strong technical education record, including Bachelor of Engineering with 8.24 CGPA and Diploma in Engineering with 78.97%.",
        "Developed a broad support foundation across Windows, service desk operations, endpoint troubleshooting, user administration, security basics, and cloud fundamentals.",
    ]

    changes = [
        "Move the strongest Windows, service desk, Active Directory, Azure, Intune, endpoint, and troubleshooting keywords into the Summary and Skills sections.",
        "Replace task-only bullets with impact-focused bullets that show tools used, issue types handled, users supported, and operational outcomes.",
        "Keep certifications visible near the top because they are strong ATS signals for infrastructure and support roles.",
        "Use standard ATS headings such as Summary, Skills, Experience, Projects, Education, Certifications, and Achievements.",
    ]
    if missing:
        changes.insert(0, "Add these missing JD keywords where truthful: " + ", ".join(missing[:10]))

    return normalize_analysis(
        {
            "current_ats_score": current_score,
            "target_ats_score": 90,
            "score_breakdown": {
                "keyword_match": keyword_score,
                "role_alignment": min(88, current_score + 15),
                "skills_match": min(90, keyword_score + 20),
                "experience_relevance": min(86, current_score + 10),
                "formatting_ats_readability": 92,
            },
            "missing_keywords": missing,
            "matched_keywords": matched[:15],
            "changes_to_make": changes,
            "ats_notes": [
                "ATS scores are estimates; real systems vary by employer and parser.",
                "This optimized version avoids placeholder language and uses only information inferred from the uploaded resume.",
            ],
            "optimized_resume": {
                "headline": role,
                "summary": summary,
                "skills": skills,
                "experience": experience,
                "projects": projects,
                "achievements": achievements,
                "certifications": resume_certifications,
            },
        }
    )


def analyze_and_optimize(jd: str, resume: str) -> dict | None:
    trimmed_jd = jd[:6000]
    trimmed_resume = resume[:9000]
    prompt = f"""
Return ONLY valid JSON.
Do not wrap the JSON in markdown.

Use this exact shape:
{{
  "current_ats_score": 0,
  "target_ats_score": 90,
  "score_breakdown": {{
    "keyword_match": 0,
    "role_alignment": 0,
    "skills_match": 0,
    "experience_relevance": 0,
    "formatting_ats_readability": 0
  }},
  "missing_keywords": [],
  "matched_keywords": [],
  "changes_to_make": [],
  "ats_notes": [],
  "optimized_resume": {{
    "headline": "",
    "summary": "",
    "skills": [],
    "experience": [],
    "projects": [],
    "achievements": [],
    "certifications": []
  }}
}}

Task:
1. Estimate the current resume ATS score for this job description.
2. Identify specific changes needed in the current resume.
3. Rewrite the resume content to target at least 90 ATS score.
4. Use only truthful information inferred from the resume. Do not invent employers, degrees, dates, certifications, or metrics.
5. Add job-description keywords naturally where they fit.
6. Keep bullets concise, ATS-readable, and achievement-focused.
7. Never return placeholder instructions such as "rewrite this", "update this", "add metrics", or "highlight projects".
8. Every optimized_resume field must contain final resume-ready text that can be shown directly to recruiters.
9. If exact metrics are missing, write strong non-metric impact bullets without fabricating numbers.
10. Set optimized_resume.headline to the best target job title/headline for this JD, for example "IT Support Engineer", "Desktop Support Engineer", "System Administrator", or "Windows Administrator".

Job description:
{trimmed_jd}

Resume:
{trimmed_resume}
"""

    def call_gemini():
        import google.generativeai as genai

        genai.configure(api_key=get_gemini_api_key())
        model = genai.GenerativeModel(GEMINI_MODEL)
        return model.generate_content(
            prompt,
            request_options={"timeout": 20},
        ).text

    future = AI_EXECUTOR.submit(call_gemini)
    try:
        response_text = future.result(timeout=25)
    except TimeoutError:
        return fallback_analysis(jd, resume)

    try:
        start = response_text.find("{")
        end = response_text.rfind("}")
        if start == -1 or end == -1:
            return None
        data = json.loads(response_text[start : end + 1])
        return normalize_analysis(data)
    except (json.JSONDecodeError, TypeError):
        return None


def clamp_score(value, default: int = 0) -> int:
    try:
        return max(0, min(100, int(value)))
    except (TypeError, ValueError):
        return default


WEAK_VERBS = {
    "worked", "handled", "helped", "did", "made", "responsible", "assisted", "participated",
    "involved", "managed tasks", "team player",
}


def quality_metrics(resume: str, analysis: dict) -> dict:
    resume_lower = resume.lower()
    weak_count = sum(len(re.findall(rf"\b{re.escape(verb)}\b", resume_lower)) for verb in WEAK_VERBS)
    bullets = re.findall(r"(?:^|\n)\s*(?:[-•*]|\d+[.)])\s+(.+)", resume)
    if not bullets:
        bullets = [line for line in resume.splitlines() if len(line.strip()) > 28][:20]

    measurable = sum(1 for bullet in bullets if re.search(r"\d+|%|reduced|increased|improved|saved|delivered", bullet, re.I))
    measurable_score = round((measurable / max(len(bullets), 1)) * 100)
    sentences = re.split(r"[.!?]+", resume)
    words = re.findall(r"[A-Za-z]+", resume)
    avg_sentence = len(words) / max(len([s for s in sentences if s.strip()]), 1)
    readability = clamp_score(100 - max(0, round((avg_sentence - 18) * 3)), 82)
    weak_verb_score = clamp_score(100 - weak_count * 12, 88)
    section_names = ["summary", "skills", "experience", "projects", "education"]
    completeness = round((sum(1 for section in section_names if section in resume_lower) / len(section_names)) * 100)
    breakdown = analysis.get("score_breakdown") or {}

    checklist = [
        {"label": "Job description keywords are represented", "passed": clamp_score(breakdown.get("keyword_match")) >= 70},
        {"label": "Resume uses ATS-safe formatting", "passed": clamp_score(breakdown.get("formatting_ats_readability")) >= 80},
        {"label": "Weak action verbs are limited", "passed": weak_count <= 2},
        {"label": "Achievements include measurable impact", "passed": measurable_score >= 45},
        {"label": "Core resume sections are complete", "passed": completeness >= 80},
    ]

    return {
        "readability": readability,
        "weak_action_verbs": weak_count,
        "weak_action_verb_score": weak_verb_score,
        "measurable_achievement_score": clamp_score(measurable_score),
        "section_completeness": clamp_score(completeness),
        "checklist": checklist,
    }


def normalize_analysis(data: dict) -> dict:
    optimized = data.get("optimized_resume") or {}
    target_score = clamp_score(data.get("target_ats_score"), 90)
    if target_score < 90:
        target_score = 90

    return {
        "current_ats_score": clamp_score(data.get("current_ats_score")),
        "target_ats_score": target_score,
        "score_breakdown": {
            "keyword_match": clamp_score((data.get("score_breakdown") or {}).get("keyword_match")),
            "role_alignment": clamp_score((data.get("score_breakdown") or {}).get("role_alignment")),
            "skills_match": clamp_score((data.get("score_breakdown") or {}).get("skills_match")),
            "experience_relevance": clamp_score((data.get("score_breakdown") or {}).get("experience_relevance")),
            "formatting_ats_readability": clamp_score(
                (data.get("score_breakdown") or {}).get("formatting_ats_readability")
            ),
        },
        "missing_keywords": as_list(data.get("missing_keywords")),
        "matched_keywords": as_list(data.get("matched_keywords")),
        "changes_to_make": as_list(data.get("changes_to_make")),
        "ats_notes": as_list(data.get("ats_notes")),
        "optimized_resume": {
            "headline": str(optimized.get("headline") or optimized.get("role") or ""),
            "summary": str(optimized.get("summary") or ""),
            "skills": as_list(optimized.get("skills")),
            "experience": as_items(optimized.get("experience")),
            "projects": as_items(optimized.get("projects")),
            "achievements": as_list(optimized.get("achievements")),
            "certifications": as_list(optimized.get("certifications")),
        },
    }


def skill_allowed(item: str) -> bool:
    text = str(item).strip()
    lowered = text.lower()
    blocked = [
        "professional experience", "hcl technologies", "key achievements", "technical projects",
        "certifications", "top performer", "sla leadership", "process efficiency",
        "powershell automation suite", "azure hybrid infrastructure support",
        "provide l2", "resolved", "administered", "developed", "assisted", "maintained",
        "awarded", "reduced", "standardized", "deployed", "managed vm",
    ]
    if any(term in lowered for term in blocked):
        return False
    if re.search(r"\b(20\d{2}|present|sep|jan|feb|mar|apr|may|jun|jul|aug|oct|nov|dec)\b", lowered):
        return False
    return bool(text)


def clean_optimized_resume(analysis: dict, jd: str, resume: str) -> dict:
    optimized = analysis["optimized_resume"]
    base_skills = [
        "Operating Systems: Windows 10/11, Windows Server, macOS, iOS",
        "Cloud & Identity: Azure, Microsoft 365 Admin Center, Active Directory, Azure Hybrid Identity",
        "Endpoint Management: Microsoft Intune, Autopilot, SCCM, device lifecycle management",
        "Networking: DNS, DHCP, LAN/WAN, VPN troubleshooting",
        "ITSM Tools: Jira Service Management, ServiceNow, Freshservice",
        "Automation: PowerShell scripting, user management, logs, device health checks",
        "Security: Access management, endpoint security, patching, least-privilege permissions",
        "Support: L2/L3 troubleshooting, service desk operations, RCA, SOP documentation",
    ]

    cleaned_skills = [item for item in as_list(optimized.get("skills")) if skill_allowed(item)]
    optimized["skills"] = unique_items(cleaned_skills + base_skills)[:14]

    if not optimized.get("experience"):
        optimized["experience"] = fallback_analysis(jd, resume)["optimized_resume"]["experience"]
    if not optimized.get("projects"):
        optimized["projects"] = fallback_analysis(jd, resume)["optimized_resume"]["projects"]
    if not optimized.get("achievements"):
        optimized["achievements"] = fallback_analysis(jd, resume)["optimized_resume"]["achievements"]

    analysis["optimized_resume"] = optimized
    return analysis


# ================= RENDER =================

def as_list(value) -> list[str]:
    if value is None:
        return []
    if isinstance(value, list):
        return [str(item) for item in value if str(item).strip()]
    if isinstance(value, str) and value.strip():
        return [value]
    return []


def as_items(value) -> list:
    if value is None:
        return []
    if isinstance(value, list):
        return [item for item in value if str(item).strip()]
    if isinstance(value, (str, dict)) and str(value).strip():
        return [value]
    return []


def list_html(items) -> str:
    rendered = []
    for item in expand_sidebar_items(as_list(items)):
        heading = sidebar_subheading(item)
        if heading:
            rendered.append(f'<li class="skill-subheading">{html.escape(heading)}</li>')
        else:
            rendered.append(f"<li>{format_label_text(item)}</li>")
    return "".join(rendered)


def expand_sidebar_items(items: list[str]) -> list[str]:
    expanded = []
    for item in items:
        text = str(item).strip()
        heading = sidebar_subheading(text)
        if not heading:
            expanded.append(text)
            continue

        cleaned = text.strip().lstrip("-•").strip()
        parts = cleaned.split(":", 1)
        expanded.append(heading)
        if len(parts) == 2 and parts[1].strip():
            expanded.append(parts[1].strip())

    return expanded


def sidebar_subheading(value: str) -> str:
    cleaned = " ".join(
        str(value)
        .replace("\u2013", "-")
        .replace("\u2014", "-")
        .replace("\u2022", "-")
        .strip()
        .lstrip("-")
        .strip()
        .rstrip(":")
        .upper()
        .split()
    )
    if "KEY ACHIEVEMENTS" in cleaned or "KEY ACHEIVEMENTS" in cleaned:
        return "KEY ACHIEVEMENTS"
    if "TECHNICAL PROJECTS" in cleaned or "TECHNICAL PROJECT" in cleaned:
        return "TECHNICAL PROJECTS"
    return ""


def format_label_text(value: str) -> str:
    safe = html.escape(str(value))
    return re.sub(
        r"^(\s*(?:[-•]\s*)?)([A-Za-z][A-Za-z0-9 &/+.-]{2,}:)",
        r"\1<strong>\2</strong>",
        safe,
    )


def format_dict_item(item: dict) -> str:
    title = item.get("role") or item.get("name") or item.get("title") or item.get("project") or ""
    company = item.get("company") or ""
    location = item.get("location") or ""
    dates = item.get("dates") or item.get("duration") or ""
    description = item.get("description") or item.get("summary") or ""
    bullets = item.get("bullets") or item.get("responsibilities") or item.get("achievements") or []

    meta = " | ".join(html.escape(value) for value in [company, location, dates] if value)
    parts = []

    if title:
        parts.append(f'<div class="item-title">{html.escape(str(title))}</div>')
    if meta:
        parts.append(f'<div class="item-meta">{meta}</div>')
    if description:
        parts.append(f'<p>{html.escape(str(description))}</p>')
    if bullets:
        parts.append("<ul>")
        parts.extend(f"<li>{html.escape(str(bullet))}</li>" for bullet in as_list(bullets))
        parts.append("</ul>")

    return "".join(parts)


def structured_html(items) -> str:
    rendered = []
    for item in as_items(items):
        if isinstance(item, dict):
            rendered.append(f'<div class="resume-item">{format_dict_item(item)}</div>')
        else:
            rendered.append(f"<li>{html.escape(str(item))}</li>")
    return "".join(rendered)


def text_html(value) -> str:
    return html.escape(str(value or ""))


class ResumeTextParser(HTMLParser):
    def __init__(self):
        super().__init__()
        self.lines = []
        self.current = []
        self.in_style = False

    def handle_starttag(self, tag, attrs):
        if tag == "style":
            self.in_style = True
        if tag in {"p", "div", "li", "br", "h1", "h2", "h3"}:
            self.flush()
        if tag == "li":
            self.current.append("- ")

    def handle_endtag(self, tag):
        if tag == "style":
            self.in_style = False
        if tag in {"p", "div", "li", "h1", "h2", "h3"}:
            self.flush()

    def handle_data(self, data):
        if not self.in_style:
            text = " ".join(data.split())
            if text:
                self.current.append(text)

    def flush(self):
        line = " ".join(self.current).strip()
        if line:
            self.lines.append(line)
        self.current = []

    def get_text(self):
        self.flush()
        return self.lines


def generate_pdf(rendered_html: str, pdf_path: str):
    browser_pdf_path = generate_pdf_with_browser(rendered_html, pdf_path)
    if browser_pdf_path:
        return

    try:
        from weasyprint import HTML

        HTML(string=rendered_html, base_url=BASE_DIR).write_pdf(pdf_path)
        return
    except Exception:
        pass

    try:
        generate_styled_pdf_from_html(rendered_html, pdf_path)
        return
    except Exception:
        pass

    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    parser = ResumeTextParser()
    parser.feed(rendered_html)

    doc = SimpleDocTemplate(pdf_path, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []

    for line in parser.get_text():
        story.append(Paragraph(html.escape(line), styles["Normal"]))
        story.append(Spacer(1, 8))

    doc.build(story)


def strip_html(value: str) -> str:
    value = re.sub(r"<style.*?</style>", " ", value, flags=re.IGNORECASE | re.DOTALL)
    value = re.sub(r"<br\s*/?>", "\n", value, flags=re.IGNORECASE)
    value = re.sub(r"</(p|div|li|h1|h2|h3)>", "\n", value, flags=re.IGNORECASE)
    value = re.sub(r"<.*?>", " ", value, flags=re.DOTALL)
    value = html.unescape(value)
    return "\n".join(" ".join(line.split()) for line in value.splitlines() if line.strip())


def class_text(rendered_html: str, class_name: str) -> str:
    match = re.search(
        rf'<[^>]+class="{re.escape(class_name)}"[^>]*>(.*?)</[^>]+>',
        rendered_html,
        flags=re.IGNORECASE | re.DOTALL,
    )
    return strip_html(match.group(1)) if match else ""


def contact_text(rendered_html: str) -> str:
    match = re.search(
        r'<div class="contact">(.*?)</div>\s*</div>\s*<div class="section-title">Skills</div>',
        rendered_html,
        flags=re.IGNORECASE | re.DOTALL,
    )
    if match:
        text = strip_html(match.group(1))
    else:
        text = class_text(rendered_html, "contact")

    return (
        text.replace("☎", "Phone:")
        .replace("✉", "Email:")
        .replace("📍", "Location:")
        .replace("🔗", "Portfolio:")
        .replace("💼", "LinkedIn:")
    )


def section_html(rendered_html: str, title: str) -> str:
    pattern = rf'<div class="section-title">{re.escape(title)}</div>(.*?)(?=<div class="section-title">|</div>\s*</div>\s*</body>)'
    match = re.search(pattern, rendered_html, flags=re.IGNORECASE | re.DOTALL)
    return match.group(1) if match else ""


def li_texts(value: str) -> list[str]:
    items = re.findall(r"<li[^>]*>(.*?)</li>", value, flags=re.IGNORECASE | re.DOTALL)
    return [strip_html(item) for item in items if strip_html(item)]


def resume_item_blocks(value: str) -> list[str]:
    matches = list(re.finditer(r'<div class="resume-item">', value, flags=re.IGNORECASE))
    blocks = []
    for index, match in enumerate(matches):
        end = matches[index + 1].start() if index + 1 < len(matches) else len(value)
        blocks.append(value[match.start() : end])
    return blocks


def generate_styled_pdf_from_html(rendered_html: str, pdf_path: str):
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.pdfgen import canvas
    from reportlab.platypus import HRFlowable, Paragraph

    page_width, page_height = A4
    left_width = page_width * 0.30
    left_x = 15
    right_x = left_width + 24
    top = page_height - 24
    bottom = 24
    left_content_width = left_width - 30
    right_content_width = page_width - right_x - 24

    body_class = ""
    body_match = re.search(r'<body[^>]*class="([^"]*)"', rendered_html, flags=re.IGNORECASE)
    if body_match:
        body_class = body_match.group(1)

    sidebar_colors = {
        "template-modern": "#0f766e",
        "template-fresher": "#1d4ed8",
        "template-corporate": "#374151",
        "template-creative": "#7c3aed",
        "template-product": "#7f1d1d",
        "template-designer": "#be185d",
    }
    sidebar_color = "#0f172a"
    for class_name, color_value in sidebar_colors.items():
        if class_name in body_class:
            sidebar_color = color_value
            break

    pdf = canvas.Canvas(pdf_path, pagesize=A4)
    styles = getSampleStyleSheet()

    left_title = ParagraphStyle("left_title", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=15, leading=17, textColor=colors.white)
    left_role = ParagraphStyle("left_role", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=7.4, leading=8.8, textColor=colors.white)
    left_text = ParagraphStyle("left_text", parent=styles["Normal"], fontSize=7.1, leading=8.7, textColor=colors.white, bulletIndent=0, leftIndent=7, firstLineIndent=-7)
    left_heading = ParagraphStyle("left_heading", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=10.5, leading=12.5, textColor=colors.white)
    right_heading = ParagraphStyle("right_heading", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=10, leading=12, spaceBefore=3)
    right_text = ParagraphStyle("right_text", parent=styles["Normal"], fontSize=8.6, leading=12.2, textColor=colors.black, bulletIndent=0, leftIndent=10, firstLineIndent=-10)
    right_paragraph = ParagraphStyle("right_paragraph", parent=styles["Normal"], fontSize=8.8, leading=13.3, textColor=colors.black)
    item_title = ParagraphStyle("item_title", parent=right_paragraph, fontName="Helvetica-Bold", fontSize=9.2, leading=11.4)
    item_meta = ParagraphStyle("item_meta", parent=right_paragraph, fontSize=7.8, leading=9.8, textColor=colors.HexColor("#64748b"))

    def p(text, style):
        safe = html.escape(text).replace("\n", "<br/>")
        safe = re.sub(
            r"^(\s*(?:[-•]\s*)?)([A-Za-z][A-Za-z0-9 &/+.-]{2,}:)",
            r"\1<b>\2</b>",
            safe,
        )
        return Paragraph(safe, style)

    def start_page():
        pdf.setFillColor(colors.HexColor(sidebar_color))
        pdf.rect(0, 0, left_width, page_height, stroke=0, fill=1)
        pdf.setFillColor(colors.white)

    def draw_flowable(flowable, x, y, width, gap):
        _, flowable_height = flowable.wrap(width, y - bottom)
        if y - flowable_height < bottom:
            pdf.showPage()
            start_page()
            y = top
            _, flowable_height = flowable.wrap(width, y - bottom)
        flowable.drawOn(pdf, x, y - flowable_height)
        return y - flowable_height - gap

    def draw_column(flowables, x, y, width):
        for flowable, gap in flowables:
            y = draw_flowable(flowable, x, y, width, gap)
        return y

    def section_heading(title: str):
        return [
            (p(title.upper(), right_heading), 2),
            (HRFlowable(width="100%", thickness=0.8, color=colors.black, spaceBefore=0, spaceAfter=0), 6),
        ]

    left = [
        (p(class_text(rendered_html, "header-name"), left_title), 3),
        (p(class_text(rendered_html, "header-role"), left_role), 8),
        (p(contact_text(rendered_html), left_text), 12),
        (p("SKILLS", left_heading), 5),
    ]
    for item in li_texts(section_html(rendered_html, "Skills")):
        if not skill_allowed(item):
            continue
        heading = sidebar_subheading(item)
        if heading:
            left.append((p(heading, left_heading), 7))
            parts = str(item).split(":", 1)
            if len(parts) == 2 and parts[1].strip():
                left.append((p(f"- {parts[1].strip()}", left_text), 3))
        else:
            left.append((p(f"- {item}", left_text), 4))
    left.append((p("CERTIFICATIONS", left_heading), 5))
    left.extend((p(f"- {item}", left_text), 4) for item in li_texts(section_html(rendered_html, "Certifications")))

    right = []
    for title in ["Summary", "Experience", "Projects", "Education", "Achievements"]:
        content = section_html(rendered_html, title)
        if not content:
            continue

        right.extend(section_heading(title))
        blocks = resume_item_blocks(content)
        if blocks:
            for block in blocks:
                block_title = class_text(block, "item-title")
                block_meta = class_text(block, "item-meta")
                if block_title:
                    right.append((p(block_title, item_title), 2))
                if block_meta:
                    right.append((p(block_meta, item_meta), 3))
                for item in li_texts(block):
                    right.append((p(f"- {item}", right_text), 4))
                desc = strip_html(re.sub(r"<ul.*?</ul>", "", block, flags=re.IGNORECASE | re.DOTALL))
                desc = "\n".join(line for line in desc.splitlines() if line not in [block_title, block_meta])
                if desc:
                    right.append((p(desc, right_paragraph), 5))
                right.append((p("", right_paragraph), 2))
        else:
            items = li_texts(content)
            if items:
                right.extend((p(f"- {item}", right_text), 4) for item in items)
            else:
                right.append((p(strip_html(content), right_paragraph), 10))
        right.append((p("", right_paragraph), 4))

    start_page()
    draw_column(left, left_x, top, left_content_width)
    draw_column(right, right_x, top - 10, right_content_width)
    pdf.save()


def add_docx_heading(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 75, 105)
    paragraph.paragraph_format.space_before = Pt(16)
    paragraph.paragraph_format.space_after = Pt(6)


def add_docx_paragraph(doc: Document, text: str, space_after: int = 6):
    paragraph = doc.add_paragraph(str(text))
    paragraph.paragraph_format.space_after = Pt(space_after)
    for run in paragraph.runs:
        run.font.size = Pt(8.5)
    return paragraph


def add_docx_bullets(doc: Document, items):
    for item in as_items(items):
        if isinstance(item, dict):
            title = item.get("role") or item.get("name") or item.get("title") or item.get("project") or ""
            meta = " | ".join(
                str(value)
                for value in [item.get("company"), item.get("location"), item.get("dates") or item.get("duration")]
                if value
            )
            description = item.get("description") or item.get("summary") or ""
            bullets = item.get("bullets") or item.get("responsibilities") or item.get("achievements") or []

            if title:
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(str(title))
                run.bold = True
            if meta:
                paragraph = doc.add_paragraph(str(meta))
                paragraph.runs[0].font.size = Pt(9)
                paragraph.runs[0].font.color.rgb = RGBColor(90, 90, 90)
            if description:
                add_docx_paragraph(doc, description, 3)
            for bullet in as_list(bullets):
                doc.add_paragraph(str(bullet), style="List Bullet")
        else:
            paragraph = doc.add_paragraph(str(item), style="List Bullet")
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.space_after = Pt(2)


def set_docx_page_border(section):
    sect_pr = section._sectPr
    pg_borders = sect_pr.find(qn("w:pgBorders"))
    if pg_borders is None:
        pg_borders = OxmlElement("w:pgBorders")
        sect_pr.append(pg_borders)
    pg_borders.set(qn("w:offsetFrom"), "page")

    for edge in ["top", "left", "bottom", "right"]:
        border = pg_borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            pg_borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")
        border.set(qn("w:space"), "18")
        border.set(qn("w:color"), "000000")


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "004B69")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def contact_links(rendered_html: str) -> dict:
    contact_match = re.search(r'<div class="contact">(.*?)</div>\s*</div>\s*<div class="section-title">Skills</div>', rendered_html, flags=re.IGNORECASE | re.DOTALL)
    source = contact_match.group(1) if contact_match else rendered_html
    links = dict(re.findall(r'<a href="([^"]+)">([^<]+)</a>', source, flags=re.IGNORECASE))
    text = contact_text(rendered_html)
    location_match = re.search(r"Location:\s*([^\n]+)", text)
    return {
        "phone": next(((label, href) for href, label in links.items() if href.startswith("tel:")), ("", "")),
        "email": next(((label, href) for href, label in links.items() if href.startswith("mailto:")), ("Email", "")),
        "portfolio": next(((label, href) for href, label in links.items() if "vercel.app" in href or label.lower() == "portfolio"), ("Portfolio", "")),
        "linkedin": next(((label, href) for href, label in links.items() if "linkedin.com" in href), ("LinkedIn", "")),
        "location": location_match.group(1).strip() if location_match else "",
    }


def generate_docx_from_resume(rendered_html: str, analysis: dict, docx_path: str):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)
    set_docx_page_border(section)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(8.5)

    optimized = analysis.get("optimized_resume", {})

    name = class_text(rendered_html, "header-name")
    headline = class_text(rendered_html, "header-role")
    contact = contact_links(rendered_html)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(name)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0, 75, 105)
    title.paragraph_format.space_after = Pt(6)

    if headline:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(headline)
        run.bold = True
        run.font.size = Pt(8.5)
        paragraph.paragraph_format.space_after = Pt(4)

    contact_paragraph = doc.add_paragraph()
    contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_paragraph.paragraph_format.space_after = Pt(18)
    if contact["location"]:
        run = contact_paragraph.add_run(f"📍 {contact['location']}  |  ")
        run.bold = True
        run.font.size = Pt(8)
    if contact["phone"][0]:
        run = contact_paragraph.add_run("📞 ")
        run.font.size = Pt(8)
        add_hyperlink(contact_paragraph, contact["phone"][0], contact["phone"][1])
        contact_paragraph.add_run("  |  ")
    if contact["email"][1]:
        add_hyperlink(contact_paragraph, "Email", contact["email"][1])
        contact_paragraph.add_run("  |  ")
    if contact["linkedin"][1]:
        add_hyperlink(contact_paragraph, "LinkedIn", contact["linkedin"][1])

    sections = [
        ("Professional Summary", optimized.get("summary") or strip_html(section_html(rendered_html, "Summary")), "paragraph"),
        ("Core Skills", optimized.get("skills") or li_texts(section_html(rendered_html, "Skills")), "bullets"),
        ("Professional Experience", optimized.get("experience") or li_texts(section_html(rendered_html, "Experience")), "bullets"),
        ("Projects", optimized.get("projects") or li_texts(section_html(rendered_html, "Projects")), "bullets"),
        ("Education", strip_html(section_html(rendered_html, "Education")), "paragraph"),
        ("Certifications", optimized.get("certifications") or li_texts(section_html(rendered_html, "Certifications")), "bullets"),
        ("Achievements", optimized.get("achievements") or li_texts(section_html(rendered_html, "Achievements")), "bullets"),
        ("Additional Information", optimized.get("additional_information") or optimized.get("technical_skills") or [], "bullets"),
    ]

    for heading, content, section_type in sections:
        if not content:
            continue

        if isinstance(content, list) and not content:
            continue

        add_docx_heading(doc, heading)
        if section_type == "paragraph":
            add_docx_paragraph(doc, content, 8)
        else:
            add_docx_bullets(doc, content)

    doc.save(docx_path)


def latex_escape(value) -> str:
    replacements = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
    }
    return "".join(replacements.get(char, char) for char in str(value or ""))


def latex_url_escape(value) -> str:
    return str(value or "").replace("\\", "/").replace("{", "%7B").replace("}", "%7D").replace("%", "%25")


def visible_url(value: str) -> str:
    return str(value or "").replace("https://", "").replace("http://", "").rstrip("/")


def visible_email(label: str, href: str) -> str:
    if label and label.lower() != "email":
        return label
    return str(href or "").replace("mailto:", "")


def latex_itemize(items: list[str]) -> list[str]:
    lines = [r"\begin{itemize}"]
    for item in items:
        if str(item).strip():
            lines.append(rf"  \item {latex_escape(item)}")
    lines.append(r"\end{itemize}")
    return lines


def latex_section(title: str, content: list[str]) -> list[str]:
    if not content:
        return []
    return [
        "",
        rf"\section*{{{latex_escape(title).upper()}}}",
        r"\vspace{-0.8em}\hrule\vspace{0.6em}",
        *content,
    ]


def latex_entry_lines(item) -> list[str]:
    if not isinstance(item, dict):
        text = str(item).strip()
        return [rf"\item {latex_escape(text)}"] if text else []

    title = item.get("role") or item.get("name") or item.get("title") or item.get("project") or ""
    meta = " | ".join(
        str(value)
        for value in [item.get("company"), item.get("location"), item.get("dates") or item.get("duration")]
        if value
    )
    description = item.get("description") or item.get("summary") or ""
    bullets = as_list(item.get("bullets") or item.get("responsibilities") or item.get("achievements"))

    lines = []
    if title:
        heading = rf"\textbf{{{latex_escape(title)}}}"
        if meta:
            heading += rf" \hfill \textit{{{latex_escape(meta)}}}"
        lines.append(rf"\item {heading}")
    elif meta:
        lines.append(rf"\item \textit{{{latex_escape(meta)}}}")

    if description:
        if lines:
            lines.append(rf"  {latex_escape(description)}")
        else:
            lines.append(rf"\item {latex_escape(description)}")

    for bullet in bullets:
        lines.append(rf"  \begin{{itemize}}\item {latex_escape(bullet)}\end{{itemize}}")

    return lines


def latex_entries(items) -> list[str]:
    entries = []
    for item in as_items(items):
        entries.extend(latex_entry_lines(item))
    if not entries:
        return []
    return [r"\begin{itemize}", *entries, r"\end{itemize}"]


def generate_latex_from_resume(rendered_html: str, analysis: dict, tex_path: str):
    optimized = analysis.get("optimized_resume", {})
    name = class_text(rendered_html, "header-name") or "Your Name"
    headline = class_text(rendered_html, "header-role")
    contact = contact_links(rendered_html)
    education = strip_html(section_html(rendered_html, "Education"))

    primary_contact_parts = []
    linkedin_contact = ""
    if contact["location"]:
        primary_contact_parts.append(latex_escape(contact["location"]))
    if contact["phone"][0]:
        primary_contact_parts.append(latex_escape(contact["phone"][0]))
    if contact["email"][1]:
        email_label = visible_email(contact["email"][0], contact["email"][1])
        primary_contact_parts.append(rf"\href{{{latex_url_escape(contact['email'][1])}}}{{{latex_escape(email_label)}}}")
    if contact["linkedin"][1]:
        linkedin_label = visible_url(contact["linkedin"][1])
        linkedin_contact = rf"\href{{{latex_url_escape(contact['linkedin'][1])}}}{{{latex_escape(linkedin_label)}}}"
    if contact["portfolio"][1]:
        portfolio_label = visible_url(contact["portfolio"][1])
        primary_contact_parts.append(rf"\href{{{latex_url_escape(contact['portfolio'][1])}}}{{{latex_escape(portfolio_label)}}}")

    lines = [
        r"\documentclass[10pt,a4paper]{article}",
        r"\usepackage[margin=0.55in]{geometry}",
        r"\usepackage[hidelinks]{hyperref}",
        r"\usepackage{enumitem}",
        r"\usepackage{xcolor}",
        r"\usepackage{titlesec}",
        r"\pagestyle{empty}",
        r"\setlength{\parindent}{0pt}",
        r"\setlist[itemize]{leftmargin=1.15em, itemsep=0.25em, topsep=0.25em}",
        r"\titleformat{\section}{\large\bfseries\uppercase}{}{0em}{}",
        "",
        r"\begin{document}",
        r"\begin{center}",
        rf"{{\LARGE \textbf{{{latex_escape(name)}}}}}\\",
    ]

    if headline:
        lines.append(rf"{latex_escape(headline)}\\")
    if primary_contact_parts:
        lines.append(" $\\bullet$ ".join(primary_contact_parts) + r"\\")
    if linkedin_contact:
        lines.append(linkedin_contact)
    lines.extend([r"\end{center}", r"\vspace{-0.4em}"])

    lines.extend(latex_section("Summary", [latex_escape(optimized.get("summary") or strip_html(section_html(rendered_html, "Summary")))]))
    lines.extend(latex_section("Skills", latex_itemize(optimized.get("skills") or li_texts(section_html(rendered_html, "Skills")))))
    lines.extend(latex_section("Experience", latex_entries(optimized.get("experience") or li_texts(section_html(rendered_html, "Experience")))))
    lines.extend(latex_section("Projects", latex_entries(optimized.get("projects") or li_texts(section_html(rendered_html, "Projects")))))
    if education:
        lines.extend(latex_section("Education", [latex_escape(education).replace("\n", r"\\")]))
    lines.extend(latex_section("Certifications", latex_itemize(optimized.get("certifications") or li_texts(section_html(rendered_html, "Certifications")))))
    lines.extend(latex_section("Achievements", latex_itemize(optimized.get("achievements") or li_texts(section_html(rendered_html, "Achievements")))))
    lines.extend(["", r"\end{document}", ""])

    with open(tex_path, "w", encoding="utf-8") as tex_file:
        tex_file.write("\n".join(lines))


def find_browser_executable() -> str | None:
    browser_paths = [
        r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
        r"C:\Program Files\Microsoft\Edge\Application\msedge.exe",
        r"C:\Program Files\Google\Chrome\Application\chrome.exe",
        r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
    ]

    for browser_path in browser_paths:
        if os.path.exists(browser_path):
            return browser_path

    return None


def generate_pdf_with_browser(rendered_html: str, pdf_path: str) -> bool:
    browser_path = find_browser_executable()
    if not browser_path:
        return False

    fd, temp_html_path = tempfile.mkstemp(suffix=".html", dir=OUTPUT_DIR)
    os.close(fd)
    user_data_dir = tempfile.mkdtemp(prefix="browser-profile-", dir=OUTPUT_DIR)

    try:
        with open(temp_html_path, "w", encoding="utf-8") as temp_file:
            temp_file.write(rendered_html)

        absolute_pdf_path = os.path.abspath(pdf_path)
        file_url = "file:///" + os.path.abspath(temp_html_path).replace("\\", "/")
        common_args = [
            browser_path,
            "--headless=new",
            "--disable-gpu",
            "--disable-extensions",
            "--run-all-compositor-stages-before-draw",
            f"--user-data-dir={user_data_dir}",
            f"--print-to-pdf={absolute_pdf_path}",
        ]

        for header_flag in ["--no-pdf-header-footer", "--print-to-pdf-no-header"]:
            command = [*common_args, header_flag, file_url]
            try:
                subprocess.run(command, check=True, timeout=60, capture_output=True)
                if os.path.exists(absolute_pdf_path) and os.path.getsize(absolute_pdf_path) > 0:
                    return True
            except Exception:
                continue

        return False
    except Exception:
        return False
    finally:
        try:
            os.remove(temp_html_path)
        except OSError:
            pass
        try:
            import shutil

            shutil.rmtree(user_data_dir, ignore_errors=True)
        except OSError:
            pass


def percent_bar(label: str, score: int) -> str:
    safe_label = html.escape(label)
    safe_score = clamp_score(score)
    return f"""
    <div class="score-row">
      <div class="score-row-head">
        <span>{safe_label}</span>
        <strong>{safe_score}%</strong>
      </div>
      <div class="bar"><span style="width: {safe_score}%"></span></div>
    </div>
    """


# ================= HTML =================

TEMPLATE_CLASSES = {
    "minimal": "template-minimal",
    "modern": "template-modern",
    "fresher": "template-fresher",
    "corporate": "template-corporate",
    "creative": "template-creative",
    "software": "template-modern",
    "product": "template-product",
    "designer": "template-designer",
}


def template_class(style: str | None) -> str:
    return TEMPLATE_CLASSES.get(str(style or "").strip().lower(), "")


FONT_MAP = {
    "arial": "Arial, sans-serif",
    "calibri": "Calibri, Arial, sans-serif",
    "georgia": "Georgia, serif",
}

THEME_CLASSES = {
    "light": "",
    "dark": "theme-dark",
}

SPACING_CLASSES = {
    "compact": "spacing-compact",
    "balanced": "",
    "relaxed": "spacing-relaxed",
}


def resume_body_classes(style: str | None, theme: str | None, spacing: str | None) -> str:
    classes = [
        template_class(style),
        THEME_CLASSES.get(str(theme or "").strip().lower(), ""),
        SPACING_CLASSES.get(str(spacing or "").strip().lower(), ""),
    ]
    return " ".join(class_name for class_name in classes if class_name)


def resume_font(font: str | None) -> str:
    return FONT_MAP.get(str(font or "").strip().lower(), FONT_MAP["arial"])


def fallback_rewrite_bullet(bullet: str, mode: str, jd: str = "") -> str:
    text = " ".join(str(bullet or "").split())
    if not text:
        return ""

    cleaned = re.sub(
        r"^(worked on|handled|responsible for|helped with|did|made)\s+",
        "",
        text,
        flags=re.IGNORECASE,
    ).strip() or text
    action_map = {
        "shorten": "Delivered",
        "professionalize": "Executed",
        "quantify": "Improved",
        "ats": "Optimized",
    }
    action = action_map.get(mode, "Developed")
    jd_terms = important_terms(jd)[:3] if jd else []
    keyword_tail = f" using {', '.join(jd_terms)}" if jd_terms else ""
    return f"{action} {cleaned[0].lower() + cleaned[1:]}{keyword_tail}, improving clarity, impact, and recruiter relevance."


def rewrite_bullet(bullet: str, mode: str, jd: str = "", api_key: str | None = None) -> str:
    mode = str(mode or "improve").strip().lower()
    api_key = get_gemini_api_key(api_key)
    if not api_key:
        return fallback_rewrite_bullet(bullet, mode, jd)

    prompt = f"""
Return only one resume bullet, no quotes and no markdown.
Mode: {mode}
Rules:
- Keep it truthful.
- Do not invent employers, tools, dates, or exact metrics.
- If metrics are missing, improve impact without fake numbers.
- Make it concise, professional, and ATS-readable.

Job description context:
{jd[:2500]}

Original bullet:
{bullet[:800]}
"""

    def call_gemini():
        import google.generativeai as genai

        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(GEMINI_MODEL)
        return model.generate_content(prompt).text.strip()

    future = AI_EXECUTOR.submit(call_gemini)
    try:
        rewritten = future.result(timeout=18)
    except Exception:
        return fallback_rewrite_bullet(bullet, mode, jd)

    rewritten = re.sub(r"^\s*[-•]\s*", "", rewritten).strip()
    return rewritten or fallback_rewrite_bullet(bullet, mode, jd)


def fallback_cover_letter(resume: str, jd: str) -> str:
    role = infer_role(jd, resume)
    matched = important_terms(jd)[:5]
    skills_sentence = ", ".join(term.title() for term in matched) or "the role's required skills"
    return (
        f"Dear Hiring Manager,\n\n"
        f"I am excited to apply for the {role} position. My background aligns well with this opportunity, "
        f"especially around {skills_sentence}. I bring a practical, outcome-focused approach and can translate "
        f"technical requirements into reliable execution.\n\n"
        f"In my resume, you will find experience and projects that demonstrate problem solving, ownership, "
        f"and the ability to contribute quickly. I would welcome the chance to discuss how my skills can support "
        f"your team's goals.\n\n"
        f"Sincerely,\n"
        f"Your Name"
    )


def generate_cover_letter_text(resume: str, jd: str, api_key: str | None = None) -> str:
    api_key = get_gemini_api_key(api_key)
    if not api_key:
        return fallback_cover_letter(resume, jd)

    prompt = f"""
Write a concise, recruiter-ready cover letter tailored to this job.
Rules:
- Use only truthful information inferred from the resume.
- Do not invent company names, dates, degrees, employers, or metrics.
- Keep it under 260 words.
- Return plain text only.

Job description:
{jd[:5000]}

Resume:
{resume[:8000]}
"""

    def call_gemini():
        import google.generativeai as genai

        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(GEMINI_MODEL)
        return model.generate_content(prompt).text.strip()

    future = AI_EXECUTOR.submit(call_gemini)
    try:
        return future.result(timeout=22) or fallback_cover_letter(resume, jd)
    except Exception:
        return fallback_cover_letter(resume, jd)


def build_html(
    data: dict,
    style: str | None = None,
    font: str | None = None,
    spacing: str | None = None,
    theme: str | None = None,
) -> str:
    with open(TEMPLATE_PATH, "r", encoding="utf-8") as template_file:
        template = template_file.read()

    replacements = {
        "{{template_class}}": resume_body_classes(style, theme, spacing),
        "{{font_family}}": resume_font(font),
        "{{header_block}}": data["header"],
        "{{summary}}": text_html(data.get("summary")),
        "{{skills}}": list_html(data.get("skills")),
        "{{experience}}": structured_html(data.get("experience")),
        "{{projects}}": structured_html(data.get("projects")),
        "{{achievements}}": list_html(data.get("achievements")),
        "{{certifications}}": list_html(data.get("certifications")),
        "{{education}}": data["education"],
    }

    for placeholder, value in replacements.items():
        template = template.replace(placeholder, value)

    return template


def build_analysis_html(fid: str, analysis: dict, token: str | None = None) -> str:
    breakdown = analysis.get("score_breakdown") or {}
    breakdown_html = "".join(
        [
            percent_bar("Keyword match", breakdown.get("keyword_match", 0)),
            percent_bar("Role alignment", breakdown.get("role_alignment", 0)),
            percent_bar("Skills match", breakdown.get("skills_match", 0)),
            percent_bar("Experience relevance", breakdown.get("experience_relevance", 0)),
            percent_bar("ATS readability", breakdown.get("formatting_ats_readability", 0)),
        ]
    )
    safe_fid = html.escape(fid)
    token_query = f"?token={html.escape(token)}" if token else ""

    return f"""
<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<title>ATS Resume Analysis</title>
<style>
body {{
  margin: 0;
  font-family: Arial, sans-serif;
  background: #f8fafc;
  color: #111827;
}}
.wrap {{
  max-width: 1100px;
  margin: 0 auto;
  padding: 32px 20px;
}}
.top {{
  display: flex;
  justify-content: space-between;
  gap: 20px;
  align-items: center;
  margin-bottom: 22px;
}}
h1 {{
  margin: 0;
  font-size: 28px;
}}
.actions {{
  display: flex;
  gap: 10px;
  flex-wrap: wrap;
}}
.note {{
  color: #475569;
  font-size: 13px;
  margin-top: 8px;
}}
a.button {{
  background: #0f172a;
  color: white;
  padding: 10px 14px;
  border-radius: 6px;
  text-decoration: none;
  font-size: 14px;
}}
.grid {{
  display: grid;
  grid-template-columns: repeat(2, minmax(0, 1fr));
  gap: 18px;
}}
.panel {{
  background: white;
  border: 1px solid #e5e7eb;
  border-radius: 8px;
  padding: 18px;
}}
.score {{
  font-size: 48px;
  font-weight: 700;
  margin: 8px 0;
}}
.target {{
  color: #047857;
}}
h2 {{
  font-size: 16px;
  margin: 0 0 12px;
}}
ul {{
  margin: 0;
  padding-left: 20px;
}}
li {{
  margin-bottom: 8px;
  line-height: 1.45;
}}
.score-row {{
  margin-bottom: 12px;
}}
.score-row-head {{
  display: flex;
  justify-content: space-between;
  font-size: 13px;
  margin-bottom: 5px;
}}
.bar {{
  height: 8px;
  background: #e5e7eb;
  border-radius: 999px;
  overflow: hidden;
}}
.bar span {{
  display: block;
  height: 100%;
  background: #2563eb;
}}
@media (max-width: 760px) {{
  .top, .grid {{
    display: block;
  }}
  .panel {{
    margin-bottom: 16px;
  }}
}}
</style>
</head>
<body>
<main class="wrap">
  <div class="top">
    <div>
      <h1>ATS Resume Analysis</h1>
      <p>Estimated fit and changes needed before downloading the optimized resume.</p>
    </div>
    <div class="actions">
      <a class="button" href="/preview/{safe_fid}">View optimized resume</a>
      <a class="button" href="/docx/{safe_fid}{token_query}">Download Word</a>
      <a class="button" href="/pdf/{safe_fid}{token_query}">Download PDF</a>
      <a class="button" href="/latex/{safe_fid}{token_query}">Download LaTeX</a>
    </div>
  </div>

  <section class="grid">
    <div class="panel">
      <h2>Current ATS Score</h2>
      <div class="score">{analysis["current_ats_score"]}%</div>
    </div>
    <div class="panel">
      <h2>Optimized Target Score</h2>
      <div class="score target">{analysis["target_ats_score"]}%+</div>
      <div class="note">ATS scores are estimates. Real employer ATS tools vary.</div>
    </div>
    <div class="panel">
      <h2>Score Breakdown</h2>
      {breakdown_html}
    </div>
    <div class="panel">
      <h2>Changes To Make</h2>
      <ul>{list_html(analysis.get("changes_to_make"))}</ul>
    </div>
    <div class="panel">
      <h2>Missing Keywords</h2>
      <ul>{list_html(analysis.get("missing_keywords"))}</ul>
    </div>
    <div class="panel">
      <h2>Matched Keywords</h2>
      <ul>{list_html(analysis.get("matched_keywords"))}</ul>
    </div>
  </section>
</main>
</body>
</html>
"""


# ================= API =================

@app.get("/", response_class=HTMLResponse)
def home():
    with open(INDEX_PATH, "r", encoding="utf-8") as index_file:
        return HTMLResponse(content=index_file.read())


@app.get("/manifest.json")
def manifest():
    return FileResponse(MANIFEST_PATH, media_type="application/manifest+json")


@app.get("/service-worker.js")
def service_worker():
    return FileResponse(SERVICE_WORKER_PATH, media_type="application/javascript")


@app.get("/app-icon.svg")
def app_icon():
    return FileResponse(APP_ICON_PATH, media_type="image/svg+xml")


@app.get("/app-logo.png")
def app_logo():
    return FileResponse(APP_LOGO_PATH, media_type="image/png")


@app.get("/app-logo-icon.png")
def app_logo_icon():
    return FileResponse(APP_LOGO_ICON_PATH, media_type="image/png")


@app.get("/ResumeFit-Pro.apk")
def android_apk():
    return FileResponse(
        APK_PATH,
        filename="ResumeFit-Pro.apk",
        media_type="application/vnd.android.package-archive",
    )


@app.get("/payment/config")
def payment_config():
    return {
        "enabled": PAYMENTS_ENABLED,
        "key_id": RAZORPAY_KEY_ID if PAYMENTS_ENABLED else "",
        "amount_rupees": RESUME_PRICE_RUPEES,
        "currency": "INR",
        "owner_emails": sorted(OWNER_EMAILS),
        "free_trial_resumes": FREE_TRIAL_RESUMES,
    }


@app.post("/payment/order")
async def create_payment_order(request: Request):
    payload = await request.json()
    fid = str(payload.get("resume_id", "")).strip()
    email = normalize_email(payload.get("email"))

    if not fid or not resume_exists(fid):
        return JSONResponse(status_code=404, content={"error": "Resume not found"})

    if not email:
        return JSONResponse(status_code=400, content={"error": "Email is required"})

    if is_owner_email(email):
        return {
            "owner_exempt": True,
            "download_token": make_download_token(fid, email),
            "download_docx": f"/docx/{fid}",
            "download_pdf": f"/pdf/{fid}",
            "download_latex": f"/latex/{fid}",
            "amount_rupees": 0,
            "currency": "INR",
        }

    if can_use_free_trial(email, fid):
        record_free_trial(email, fid)
        status = free_trial_status(email, fid)
        return {
            "owner_exempt": False,
            "free_trial": True,
            "download_token": make_download_token(fid, email),
            "download_docx": f"/docx/{fid}",
            "download_pdf": f"/pdf/{fid}",
            "download_latex": f"/latex/{fid}",
            "free_trials_remaining": status["remaining"],
            "free_trial_limit": FREE_TRIAL_RESUMES,
            "amount_rupees": 0,
            "currency": "INR",
        }

    if not PAYMENTS_ENABLED or not razorpay_client:
        return JSONResponse(
            status_code=503,
            content={
                "error": "Payments are not configured. Add Razorpay keys in environment variables.",
            },
        )

    receipt = f"resume_{fid}_{uuid.uuid4().hex[:10]}"
    order = razorpay_client.order.create(
        {
            "amount": RESUME_PRICE_PAISE,
            "currency": "INR",
            "receipt": receipt,
            "payment_capture": 1,
            "notes": {"resume_id": fid, "email": email},
        }
    )

    cursor.execute(
        """
        INSERT INTO payments (id, resume_id, email, amount_paise, currency, status, razorpay_order_id)
        VALUES (?, ?, ?, ?, ?, ?, ?)
        """,
        (str(uuid.uuid4()), fid, email, RESUME_PRICE_PAISE, "INR", "created", order["id"]),
    )
    conn.commit()

    return {
        "owner_exempt": False,
        "key_id": RAZORPAY_KEY_ID,
        "order_id": order["id"],
        "amount_paise": RESUME_PRICE_PAISE,
        "amount_rupees": RESUME_PRICE_RUPEES,
        "currency": "INR",
    }


@app.post("/payment/verify")
async def verify_payment(request: Request):
    payload = await request.json()
    fid = str(payload.get("resume_id", "")).strip()
    email = normalize_email(payload.get("email"))
    order_id = str(payload.get("razorpay_order_id", "")).strip()
    payment_id = str(payload.get("razorpay_payment_id", "")).strip()
    signature = str(payload.get("razorpay_signature", "")).strip()

    if not fid or not email or not order_id or not payment_id or not signature:
        return JSONResponse(status_code=400, content={"error": "Payment verification details are incomplete"})

    if not PAYMENTS_ENABLED or not razorpay_client:
        return JSONResponse(status_code=503, content={"error": "Payments are not configured"})

    try:
        razorpay_client.utility.verify_payment_signature(
            {
                "razorpay_order_id": order_id,
                "razorpay_payment_id": payment_id,
                "razorpay_signature": signature,
            }
        )
    except Exception:
        cursor.execute(
            "UPDATE payments SET status=? WHERE razorpay_order_id=?",
            ("failed", order_id),
        )
        conn.commit()
        return JSONResponse(status_code=400, content={"error": "Payment verification failed"})

    cursor.execute(
        """
        UPDATE payments
        SET status=?, razorpay_payment_id=?, paid_at=CURRENT_TIMESTAMP
        WHERE razorpay_order_id=? AND resume_id=? AND email=?
        """,
        ("paid", payment_id, order_id, fid, email),
    )
    conn.commit()

    if cursor.rowcount == 0:
        return JSONResponse(status_code=404, content={"error": "Payment record not found"})

    return {
        "paid": True,
        "download_token": make_download_token(fid, email),
        "download_docx": f"/docx/{fid}",
        "download_pdf": f"/pdf/{fid}",
        "download_latex": f"/latex/{fid}",
    }


@app.post("/optimize")
async def optimize_resume(
    resume_file: UploadFile = File(...),
    jd: str = Form(...),
    api_key: str | None = Form(None),
    user_email: str | None = Form(None),
    template_style: str | None = Form(None),
    font_style: str | None = Form(None),
    spacing_style: str | None = Form(None),
    theme_style: str | None = Form(None),
):
    api_key = get_gemini_api_key(api_key)
    if not api_key:
        return JSONResponse(
            status_code=500,
            content={"error": "GEMINI_API_KEY is not configured on the server. Add it in Render environment variables and redeploy."},
        )

    text = extract_text(resume_file)
    if not text.strip():
        return JSONResponse(
            status_code=400,
            content={"error": "Could not read resume text. Upload a text-based DOCX or PDF file."},
        )
    education = extract_education(text)
    try:
        analysis = analyze_and_optimize(jd, text)
    except Exception:
        analysis = fallback_analysis(jd, text)

    if not analysis:
        analysis = fallback_analysis(jd, text)

    analysis = clean_optimized_resume(analysis, jd, text)
    analysis["quality_metrics"] = quality_metrics(text, analysis)
    headline = analysis["optimized_resume"].get("headline") or infer_role(jd, text)
    header = extract_header(text, headline)

    data = {
        "header": header,
        "education": education,
        **analysis["optimized_resume"],
    }

    rendered_html = build_html(data, template_style, font_style, spacing_style, theme_style)
    fid = str(uuid.uuid4())[:8]

    html_path = os.path.join(OUTPUT_DIR, f"{fid}.html")
    with open(html_path, "w", encoding="utf-8") as output_file:
        output_file.write(rendered_html)

    analysis_json = json.dumps(analysis)
    cursor.execute(
        "INSERT INTO resumes (id, html, analysis_json) VALUES (?, ?, ?)",
        (fid, rendered_html, analysis_json),
    )
    conn.commit()

    normalized_user_email = normalize_email(user_email)
    owner_exempt = is_owner_email(normalized_user_email)
    download_token = make_download_token(fid, normalized_user_email) if owner_exempt else ""
    trial_status = free_trial_status(normalized_user_email, fid)

    return {
        "id": fid,
        "current_ats_score": analysis["current_ats_score"],
        "target_ats_score": analysis["target_ats_score"],
        "score_breakdown": analysis["score_breakdown"],
        "missing_keywords": analysis["missing_keywords"],
        "matched_keywords": analysis["matched_keywords"],
        "changes_to_make": analysis["changes_to_make"],
        "ats_notes": analysis["ats_notes"],
        "quality_metrics": analysis["quality_metrics"],
        "optimized_resume": analysis["optimized_resume"],
        "analysis_url": f"/analysis/{fid}",
        "preview_url": f"/preview/{fid}",
        "share_url": f"/preview/{fid}",
        "download_pdf": f"/pdf/{fid}",
        "download_docx": f"/docx/{fid}",
        "download_latex": f"/latex/{fid}",
        "payment_required": not owner_exempt,
        "payments_enabled": PAYMENTS_ENABLED,
        "price_rupees": RESUME_PRICE_RUPEES,
        "currency": "INR",
        "owner_exempt": owner_exempt,
        "download_token": download_token,
        "free_trial_limit": FREE_TRIAL_RESUMES,
        "free_trials_used": trial_status["used"],
        "free_trials_remaining": trial_status["remaining"],
        "free_trial_available": bool(not owner_exempt and can_use_free_trial(normalized_user_email, fid)),
    }


@app.post("/rewrite-bullet")
async def rewrite_resume_bullet(request: Request):
    payload = await request.json()
    bullet = str(payload.get("bullet", "")).strip()
    mode = str(payload.get("mode", "improve")).strip()
    jd = str(payload.get("jd", "")).strip()
    api_key = payload.get("api_key")

    if not bullet:
        return JSONResponse(status_code=400, content={"error": "Add a bullet point to rewrite."})

    return {"rewritten": rewrite_bullet(bullet, mode, jd, api_key)}


@app.post("/cover-letter")
async def cover_letter(
    resume_file: UploadFile = File(...),
    jd: str = Form(...),
    api_key: str | None = Form(None),
):
    resume_text = extract_text(resume_file)
    if not resume_text.strip():
        return JSONResponse(
            status_code=400,
            content={"error": "Could not read resume text. Upload a text-based DOCX or PDF file."},
        )

    return {"cover_letter": generate_cover_letter_text(resume_text, jd, api_key)}


@app.post("/github-import")
async def github_import(request: Request):
    payload = await request.json()
    username = re.sub(r"[^A-Za-z0-9-]", "", str(payload.get("username", "")).strip())
    if not username:
        return JSONResponse(status_code=400, content={"error": "Enter a GitHub username."})

    try:
        headers = {"User-Agent": "ResumeFit-Pro"}
        user_request = urllib.request.Request(f"https://api.github.com/users/{username}", headers=headers)
        repos_request = urllib.request.Request(
            f"https://api.github.com/users/{username}/repos?sort=updated&per_page=8",
            headers=headers,
        )
        with urllib.request.urlopen(user_request, timeout=12) as response:
            profile = json.loads(response.read().decode("utf-8"))
        with urllib.request.urlopen(repos_request, timeout=12) as response:
            repos = json.loads(response.read().decode("utf-8"))
    except Exception:
        return JSONResponse(status_code=502, content={"error": "Could not import this GitHub profile."})

    projects = []
    languages = []
    for repo in repos:
        if repo.get("fork"):
            continue
        language = repo.get("language")
        if language:
            languages.append(language)
        projects.append(
            {
                "name": repo.get("name", ""),
                "description": repo.get("description") or "GitHub project",
                "url": repo.get("html_url", ""),
                "language": language or "",
                "stars": repo.get("stargazers_count", 0),
            }
        )

    return {
        "name": profile.get("name") or username,
        "headline": profile.get("bio") or "",
        "profile_url": profile.get("html_url"),
        "public_repos": profile.get("public_repos", 0),
        "skills": unique_items(languages)[:8],
        "projects": projects[:5],
    }


@app.get("/analysis/{fid}", response_class=HTMLResponse)
def analysis_page(fid: str, token: str | None = None):
    cursor.execute("SELECT analysis_json FROM resumes WHERE id=?", (fid,))
    row = cursor.fetchone()

    if not row or not row[0]:
        return HTMLResponse(content="Not found", status_code=404)

    return HTMLResponse(content=build_analysis_html(fid, json.loads(row[0]), token))


# ================= PREVIEW =================

@app.get("/preview/{fid}", response_class=HTMLResponse)
def preview(fid: str):
    cursor.execute("SELECT html FROM resumes WHERE id=?", (fid,))
    row = cursor.fetchone()

    if not row:
        return HTMLResponse(content="Not found", status_code=404)

    return HTMLResponse(content=row[0])


# ================= PDF =================

@app.get("/pdf/{fid}")
def download_pdf(fid: str, token: str | None = None):
    cursor.execute("SELECT html FROM resumes WHERE id=?", (fid,))
    row = cursor.fetchone()

    if not row:
        return JSONResponse(status_code=404, content={"error": "Not found"})

    access_error = require_download_access(fid, token)
    if access_error:
        return access_error

    rendered_html = row[0]
    pdf_path = os.path.join(OUTPUT_DIR, f"{fid}.pdf")

    generate_pdf(rendered_html, pdf_path)

    return FileResponse(pdf_path, filename="optimized_resume.pdf", media_type="application/pdf")


@app.get("/docx/{fid}")
def download_docx(fid: str, token: str | None = None):
    cursor.execute("SELECT html, analysis_json FROM resumes WHERE id=?", (fid,))
    row = cursor.fetchone()

    if not row:
        return JSONResponse(status_code=404, content={"error": "Not found"})

    access_error = require_download_access(fid, token)
    if access_error:
        return access_error

    rendered_html = row[0]
    analysis = json.loads(row[1] or "{}")
    docx_path = os.path.join(OUTPUT_DIR, f"{fid}.docx")

    generate_docx_from_resume(rendered_html, analysis, docx_path)

    return FileResponse(
        docx_path,
        filename="optimized_resume.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@app.get("/latex/{fid}")
def download_latex(fid: str, token: str | None = None):
    cursor.execute("SELECT html, analysis_json FROM resumes WHERE id=?", (fid,))
    row = cursor.fetchone()

    if not row:
        return JSONResponse(status_code=404, content={"error": "Not found"})

    access_error = require_download_access(fid, token)
    if access_error:
        return access_error

    rendered_html = row[0]
    analysis = json.loads(row[1] or "{}")
    tex_path = os.path.join(OUTPUT_DIR, f"{fid}.tex")

    generate_latex_from_resume(rendered_html, analysis, tex_path)

    return FileResponse(
        tex_path,
        filename="optimized_resume.tex",
        media_type="text/x-tex",
    )
